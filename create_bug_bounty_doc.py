import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, qn('http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'), is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

# Create a new Document
doc = Document()
doc.add_heading('Bug Bounty Programs Offering Substantial Rewards', 0)

# List of companies and their bug bounty programs
companies = [
    ("Google", "Google Vulnerability Reward Program (VRP)", "https://bughunters.google.com/"),
    ("Microsoft", "Microsoft Bug Bounty Program", "https://www.microsoft.com/en-us/msrc/bounty"),
    ("Apple", "Apple Security Bounty", "https://developer.apple.com/security-bounty/"),
    ("Facebook (Meta)", "Meta Bug Bounty Program", "https://www.facebook.com/whitehat"),
    ("Amazon", "Amazon Vulnerability Research Program", "https://aws.amazon.com/security/vulnerability-reporting/"),
    ("GitHub", "GitHub Security Bug Bounty", "https://bounty.github.com/"),
    ("Intel", "Intel Bug Bounty Program", "https://www.intel.com/content/www/us/en/security-center/bug-bounty-program.html"),
    ("Uber", "Uber Bug Bounty Program", "https://hackerone.com/uber"),
    ("Twitter", "Twitter Bug Bounty Program", "https://hackerone.com/twitter"),
    ("Dropbox", "Dropbox Bug Bounty Program", "https://www.dropbox.com/security/bug-bounty"),
    ("PayPal", "PayPal Bug Bounty Program", "https://www.paypal.com/us/webapps/mpp/security/reporting-security-issues"),
    ("Yahoo", "Yahoo Bug Bounty Program", "https://hackerone.com/yahoo"),
    ("Slack", "Slack Bug Bounty Program", "https://hackerone.com/slack"),
    ("Shopify", "Shopify Bug Bounty Program", "https://hackerone.com/shopify"),
    ("Netflix", "Netflix Bug Bounty Program", "https://bugcrowd.com/netflix"),
    ("Adobe", "Adobe Vulnerability Disclosure Program", "https://hackerone.com/adobe"),
    ("Tesla", "Tesla Bug Bounty Program", "https://www.tesla.com/security"),
    ("Spotify", "Spotify Bug Bounty Program", "https://hackerone.com/spotify"),
    ("LinkedIn", "LinkedIn Bug Bounty Program", "https://hackerone.com/linkedin"),
    ("Oracle", "Oracle Vulnerability Reward Program", "https://www.oracle.com/security-alerts/"),
]

for company, program, url in companies:
    doc.add_heading(company, level=1)
    p = doc.add_paragraph()
    p.add_run(program).bold = True
    p.add_run(" - ")
    add_hyperlink(p, url, url)

# Determine the Downloads folder path
downloads_path = str(Path.home() / "Downloads" / "Bug_Bounty_Programs.docx")

# Save the document
doc.save(downloads_path)

print(f"Document saved to {downloads_path}")
